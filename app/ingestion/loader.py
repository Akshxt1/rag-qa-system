import os
from pathlib import Path
from typing import List, Dict, Optional
from loguru import logger

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}


class DocumentLoader:
    def load_file(self, file_path: str) -> Dict:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type '{ext}'")

        loaders = {
            ".pdf": self._load_pdf,
            ".docx": self._load_docx,
            ".txt": self._load_txt,
            ".md": self._load_txt,
        }
        return loaders[ext](path)

    def _load_pdf(self, path: Path) -> Dict:
        if pdfplumber is None:
            raise ImportError("pdfplumber not installed.")
        pages_text = []
        page_count = 0
        with pdfplumber.open(path) as pdf:
            page_count = len(pdf.pages)
            for i, page in enumerate(pdf.pages, 1):
                text = page.extract_text(x_tolerance=3, y_tolerance=3)
                if not text or not text.strip():
                    text = page.extract_text(x_tolerance=7, y_tolerance=7)
                if text and text.strip():
                    cleaned = self._clean_pdf_text(text.strip())
                    pages_text.append(f"[PAGE {i}]\n{cleaned}")

        full_text = "\n\n".join(pages_text)
        doc = self._make_doc(full_text, path, "pdf", page_count)
        doc["page_count"] = page_count
        return doc

    @staticmethod
    def _clean_pdf_text(text: str) -> str:
        import re
        words = text.split()
        if not words:
            return text
        short_ratio = sum(1 for w in words if len(w) <= 2) / len(words)
        if short_ratio < 0.45:
            return text
        # Fix doubled-char artifact: "r rr" → "r", "aa aaa" → "aa"
        text = re.sub(r'\b(\w{1,4})\s+\1{1,3}\b', r'\1', text)
        # Collapse multiple spaces
        text = re.sub(r'  +', ' ', text)
        return text.strip()

    def _load_docx(self, path: Path) -> Dict:
        if DocxDocument is None:
            raise ImportError("python-docx not installed.")
        doc = DocxDocument(path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)
        return self._make_doc("\n\n".join(paragraphs), path, "docx")

    def _load_txt(self, path: Path) -> Dict:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            full_text = f.read().strip()
        doc_type = "md" if path.suffix.lower() == ".md" else "txt"
        return self._make_doc(full_text, path, doc_type)

    @staticmethod
    def _make_doc(text, path, doc_type, page_count=None):
        return {
            "text": text,
            "source": str(path),
            "filename": path.name,
            "doc_type": doc_type,
            "page_count": page_count,
            "char_count": len(text),
        }

    def load_directory(self, dir_path: str) -> List[Dict]:
        docs = []
        for file in sorted(Path(dir_path).rglob("*")):
            if file.is_file() and file.suffix.lower() in SUPPORTED_EXTENSIONS:
                try:
                    docs.append(self.load_file(str(file)))
                except Exception as e:
                    logger.warning(f"Skipping {file.name}: {e}")
        return docs

    def load_uploaded_file(self, file_bytes: bytes, filename: str) -> Dict:
        import tempfile
        suffix = Path(filename).suffix.lower()
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        try:
            doc = self.load_file(tmp_path)
            doc["source"] = filename
            doc["filename"] = filename
            return doc
        finally:
            os.remove(tmp_path)